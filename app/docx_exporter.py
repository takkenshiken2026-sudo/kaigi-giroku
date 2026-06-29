from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt

DEFAULT_FONT = "Yu Gothic"
HEADING_FONT = "Yu Gothic"


def _set_run_font(run, size_pt: float | None = None, bold: bool = False) -> None:
    run.font.name = DEFAULT_FONT
    run.bold = bold
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), DEFAULT_FONT)


def _strip_inline_markdown(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _add_inline_formatted(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run)


def _is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    return all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _add_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_heading(_strip_inline_markdown(text), level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in heading.runs:
        _set_run_font(
            run,
            size_pt=16 if level == 1 else 14 if level == 2 else 12,
            bold=True,
        )


def _style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = DEFAULT_FONT
    normal.font.size = Pt(11)
    r_pr = normal._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), DEFAULT_FONT)


def minutes_to_docx(minutes: str) -> bytes:
    doc = Document()
    _style_document(doc)

    lines = minutes.replace("\r\n", "\n").split("\n")
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("##"):
            _add_heading(doc, stripped[2:].strip(), 1)
            index += 1
            continue

        if stripped.startswith("## "):
            _add_heading(doc, stripped[3:].strip(), 2)
            index += 1
            continue

        if stripped.startswith("### "):
            _add_heading(doc, stripped[4:].strip(), 3)
            index += 1
            continue

        if stripped.startswith("|") and "|" in stripped:
            table_lines: list[str] = []
            while index < len(lines) and "|" in lines[index]:
                table_lines.append(lines[index])
                index += 1

            rows = [
                _parse_table_row(row_line)
                for row_line in table_lines
                if not _is_table_separator(row_line)
            ]
            if rows:
                column_count = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=column_count)
                table.style = "Table Grid"
                for row_index, row_cells in enumerate(rows):
                    for col_index in range(column_count):
                        cell_text = row_cells[col_index] if col_index < len(row_cells) else ""
                        cell = table.rows[row_index].cells[col_index]
                        cell.text = ""
                        paragraph = cell.paragraphs[0]
                        _add_inline_formatted(paragraph, _strip_inline_markdown(cell_text))
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_inline_formatted(paragraph, stripped[2:])
            index += 1
            continue

        paragraph = doc.add_paragraph()
        _add_inline_formatted(paragraph, stripped)
        index += 1

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
